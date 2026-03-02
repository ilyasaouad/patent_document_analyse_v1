"""
Utility to extract text from images embedded in DOCX files using OCR.

DOCX files can contain inline images (diagrams, flowcharts, tables as images)
whose text content is invisible to python-docx's paragraph text extraction.
This module extracts those images and runs Tesseract OCR on them.
"""

import io
import os
import shutil
import subprocess
from typing import List, Tuple, Optional


def _find_tesseract() -> Optional[str]:
    """Locate the Tesseract executable on this system."""
    cmd = shutil.which("tesseract")
    if cmd:
        return cmd

    if os.name == "nt":
        candidates = [
            os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"),
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in candidates:
            if os.path.exists(path):
                return path

    return None


def _tesseract_available() -> bool:
    """Check whether Tesseract OCR engine is usable."""
    cmd = _find_tesseract()
    if not cmd:
        return False
    try:
        subprocess.run([cmd, "--version"], capture_output=True, check=True)
        return True
    except Exception:
        return False


def extract_images_with_positions(doc) -> List[Tuple[int, bytes, str]]:
    """Walk through every paragraph in the DOCX and yield (index, image_bytes, name)."""
    NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    results: List[Tuple[int, bytes, str]] = []

    for para_idx, para in enumerate(doc.paragraphs):
        blips = para._element.findall(f".//{{{NS_A}}}blip")
        for blip in blips:
            embed_id = blip.get(f"{{{NS_R}}}embed")
            if not embed_id:
                continue
            try:
                image_part = doc.part.rels[embed_id].target_part
                image_name = os.path.basename(image_part.partname)
                results.append((para_idx, image_part.blob, image_name))
            except (KeyError, AttributeError):
                continue

    return results


def ocr_image_bytes(image_bytes: bytes, lang: str = "eng") -> str:
    """Run Tesseract OCR on raw image bytes."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    tesseract_cmd = _find_tesseract()
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    try:
        # Map MinerU's 'en' language code to Tesseract's 'eng' language code
        tesseract_lang = "eng" if lang == "en" else lang
        
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode not in ("L", "RGB"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img, lang=tesseract_lang, config="--psm 6")
        return text.strip()
    except Exception as e:
        print(f"[docx_image_ocr] OCR failed: {e}")
        return ""


def _find_numbering_xml(part):
    """Find the numbering.xml part from related parts."""
    for rp in part.related_parts.values():
        if rp.partname == "/word/numbering.xml":
            return rp
    return None


def _get_para_number_text(para, num_counters: dict) -> str:
    """Get paragraph numbering text (e.g., '1.', '2.', 'a)', etc.)."""
    try:
        numPr = para._element.xpath(".//w:numPr")
        if not numPr:
            return ""
        
        numId = numPr[0].xpath(".//w:numId/@w:val")
        ilvl = numPr[0].xpath(".//w:ilvl/@w:val")
        
        if not numId or not ilvl:
            return ""
        
        numId = str(numId[0])
        ilvl = str(ilvl[0])
        key = f"{numId}_{ilvl}"
        
        num_counters[key] = num_counters.get(key, 0) + 1
        num = num_counters[key]
        
        part = para.part
        numPart = _find_numbering_xml(part)
        if not numPart:
            return f"{num}."
        
        from lxml import etree
        root = etree.fromstring(numPart.blob)
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        
        num_elems = root.xpath(f".//w:num[@w:numId='{numId}']", namespaces=nsmap)
        if not num_elems:
            return f"{num}."
        
        abstractNumId = num_elems[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId")
        
        override_xpath = f".//w:num[@w:numId='{numId}']//w:lvlOverride[@w:ilvl='{ilvl}']//w:startOverride"
        override = root.xpath(override_xpath, namespaces=nsmap)
        if override:
            start_val = override[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
            if start_val:
                return str(int(start_val) + num - 1) + "."
        
        abstract = root.xpath(f".//w:abstractNum[@w:abstractNumId='{abstractNumId}']", namespaces=nsmap)
        if not abstract:
            return f"{num}."
        
        lvl = abstract[0].xpath(f".//w:lvl[@w:ilvl='{ilvl}']", namespaces=nsmap)
        if not lvl:
            return f"{num}."
        
        start = lvl[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start")
        if start:
            return str(int(start) + num - 1) + "."
        
        fmt = lvl[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt")
        
        if fmt == "decimal":
            return f"{num}."
        elif fmt == "lowerLetter":
            return chr(96 + num) + ")"
        elif fmt == "upperLetter":
            return chr(64 + num) + ")"
        elif fmt == "lowerRoman":
            return _int_to_roman(num) + "."
        elif fmt == "upperRoman":
            return _int_to_roman(num).upper() + "."
        
        return f"{num}."
    except Exception as e:
        return ""


def _int_to_roman(num: int) -> str:
    """Convert integer to Roman numeral."""
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ["m", "cm", "d", "cd", "c", "xc", "l", "xl", "x", "ix", "v", "iv", "i"]
    result = ""
    for i in range(len(val)):
        while num >= val[i]:
            result += syms[i]
            num -= val[i]
    return result


def process_docx_with_ocr(path, lang: str = "eng") -> str:
    """Extract text from DOCX, including OCR of embedded images."""
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed."

    doc = Document(str(path))
    has_ocr = _tesseract_available()

    if not has_ocr:
        print("[docx_image_ocr] Tesseract not found – images will be skipped.")

    image_map: dict[int, List[Tuple[bytes, str]]] = {}
    if has_ocr:
        for para_idx, img_bytes, img_name in extract_images_with_positions(doc):
            image_map.setdefault(para_idx, []).append((img_bytes, img_name))

    parts: List[str] = []
    num_counters: dict = {}

    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        num_prefix = _get_para_number_text(para, num_counters)
        if para_text:
            full_text = f"{num_prefix} {para_text}" if num_prefix else para_text
            parts.append(full_text)

        if para_idx in image_map:
            for img_bytes, img_name in image_map[para_idx]:
                ocr_text = ocr_image_bytes(img_bytes, lang=lang)
                if ocr_text:
                    parts.append(f"[Image OCR — {img_name}]:\n{ocr_text}")

    return "\n\n".join(parts)
